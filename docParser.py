from docx import Document

"""returns the full text inside the word doc"""
def read_docx_text(pathToFile):
    try:
        doc = Document(pathToFile)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except FileNotFoundError:
        print("Error: File not found.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

"""
creates a cover letter by using the cover letter template and replacing the placeholders inside it
"""
def create_cover_letter(template_path, output_path, cover_letter_body, replacement_values):
    try:
        doc = Document(template_path)

        placeholders = {
            "<content>": cover_letter_body,
            "<company name>": replacement_values[0],
            "<location>": replacement_values[1],
            "<current date>": replacement_values[2]
        }
        # Replace placeholders
        for para in doc.paragraphs:
            for key, value in placeholders.items():
                if key in para.text and value != "NULL":
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        # Save the document safely as a new file
        doc.save(output_path)
        print(f"✅ Cover letter saved at {output_path}")

    except Exception as e:
        print(f"❌ Error: {e}")

"""returns the text from a text file"""
def extract_text(filePath):
    try:
        with open(filePath, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        print("Error: File not found.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")